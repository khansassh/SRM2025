import os
import io
import csv
import json
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, session
import psycopg2
import psycopg2.extras
from werkzeug.security import generate_password_hash, check_password_hash
from reportlab.pdfgen import canvas
from docx import Document
from datetime import datetime
import os
from datetime import datetime
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO
import logging
import pdfkit
from docx import Document
import pandas as pd

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Necessary for flash messages

app.config['REPORT_FOLDER'] = "reports"

if not os.path.exists(app.config['REPORT_FOLDER']):
    os.makedirs(app.config['REPORT_FOLDER'])

# Logging setup
logging.basicConfig(level=logging.DEBUG)

# Database connection setup
# Database connection details
DB_HOST = "localhost"
DB_NAME = "khalza"
DB_USER = "postgres"
DB_PASS = "ilovemoon19"

PDFKIT_CONFIG = pdfkit.configuration(wkhtmltopdf=r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe")

# Function to get database connection
def get_db_connection():
    conn = psycopg2.connect(
        host=DB_HOST,
        database=DB_NAME,
        user=DB_USER,
        password=DB_PASS
    )
    cursor = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    return conn, cursor

# Fetch data from PostgreSQL
def fetch_data():
    conn, cur = get_db_connection()  # Use the function instead of DB_CONFIG
    
    # Fetch assets
    cur.execute("SELECT * FROM assets;")
    assets = cur.fetchall()
    
    # Fetch risks
    cur.execute("SELECT * FROM risks;")
    risks = cur.fetchall()

    # Fetch threats
    cur.execute("SELECT * FROM threats;")
    threats = cur.fetchall()

    conn.close()
    return assets, risks, threats

def log_activity(user_id, description):
    conn, cursor = get_db_connection()
    cursor.execute(
        "INSERT INTO activity_logs (user_id, description) VALUES (%s, %s)",
        (user_id, description)
    )
    conn.commit()
    cursor.close()
    conn.close()
# Routes

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        
        # Hash the password before storing it in the database
        password_hash = generate_password_hash(password)
        
        conn, cursor = get_db_connection()
        
        # Insert the user into the database
        cursor.execute("INSERT INTO users (username, email, password_hash) VALUES (%s, %s, %s)",
                       (username, email, password_hash))
        
        conn.commit()
        cursor.close()
        conn.close()

        # In both your dashboard and dashboard_data routes, change the asset query to:
        cursor.execute("""
            SELECT 
                COUNT(*) as total,
                COALESCE(SUM(CASE WHEN critical THEN 1 ELSE 0 END), 0) as critical 
            FROM assets
        """)
        
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn, cursor = get_db_connection()
        
        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        
        cursor.close()
        conn.close()
        
        if user and check_password_hash(user['password_hash'], password):
            # ✅ Store user data in session
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['email'] = user['email']
            
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))  # Redirect to dashboard
        else:
            flash('Invalid credentials. Please try again.', 'error')
    
    return render_template('login.html')



@app.route('/api/dashboard-data')
def dashboard_data():
    conn, cursor = get_db_connection()
    
    try:
        # Risk data
        cursor.execute("SELECT risk_type, COUNT(*) as count FROM risks GROUP BY risk_type")
        risk_data = cursor.fetchall()

        # Threats
        cursor.execute("SELECT name, severity, created_at FROM threats ORDER BY created_at DESC LIMIT 5")
        threats = cursor.fetchall()

        # Assets
        cursor.execute("""
            SELECT 
                COUNT(*) as total,
                SUM(CASE WHEN critical THEN 1 ELSE 0 END) as critical 
            FROM assets
        """)
        assets = cursor.fetchone()

        return jsonify({
            'risk_data': [dict(r) for r in risk_data],
            'threats': [dict(t) for t in threats],
            'assets': dict(assets)
        })
    finally:
        cursor.close()
        conn.close()

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        flash('You must log in first.', 'error')
        return redirect(url_for('login'))

    response = get_activity_log()
    activities = response.json

    conn, cursor = get_db_connection()
    cursor.execute("SELECT * FROM users WHERE id = %s", (session['user_id'],))
    user = cursor.fetchone()

    cursor.execute("SELECT risk_type, severity, COUNT(*) as count FROM risks GROUP BY risk_type, severity")
    risk_data = [{"risk_type": row[0], "severity": row[1], "count": row[2]} for row in cursor.fetchall()]

    cursor.execute("SELECT COUNT(*) as total, SUM(CASE WHEN critical THEN 1 ELSE 0 END) as critical FROM assets")
    assets = cursor.fetchone()

    cursor.execute("SELECT * FROM threats ORDER BY created_at DESC LIMIT 5")
    threats = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'dashboard.html',
        user=user,
        activities=activities,
        risk_data=risk_data,
        threats=threats,
        assets=assets
    )


@app.route('/logout')
def logout():
    session.clear()  # Clears all session data
    flash("You have been logged out.", "info")
    return redirect(url_for('login'))

 
@app.route('/get-activity-log')
def get_activity_log():
    if 'user_id' not in session:
        return jsonify([])  # Return empty list if not logged in

    user_id = session['user_id']  # ✅ Use logged-in user ID
    conn, cursor = get_db_connection()
    
    cursor.execute(
        "SELECT timestamp, description FROM activity_logs WHERE user_id = %s ORDER BY timestamp DESC LIMIT 10",
        (user_id,)
    )
    
    activities = cursor.fetchall()
    
    cursor.close()
    conn.close()

    # Convert tuples to dictionaries
    activity_list = [{'timestamp': activity['timestamp'].strftime('%Y-%m-%d %H:%M'), 'description': activity['description']} for activity in activities]

    return jsonify([
        {"timestamp": "2025-02-09 12:30:00", "description": "Logged in"},
        {"timestamp": "2025-02-09 12:45:00", "description": "Viewed risk report"}
    ])



@app.route('/download-activity-log')
def download_activity_log():
    user_id = 1  # Replace this with the actual user ID
    conn, cursor = get_db_connection()
    cursor.execute(
        "SELECT timestamp, description FROM activity_logs WHERE user_id = %s ORDER BY timestamp DESC",
        (user_id,)
    )
    activities = cursor.fetchall()
    cursor.close()
    conn.close()

    # Generate a CSV file
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Timestamp', 'Description'])
    for activity in activities:
        writer.writerow([activity[0], activity[1]])
    output.seek(0)

    # Use 'as_attachment=True' and specify 'filename'
    return send_file(io.BytesIO(output.getvalue().encode()), as_attachment=True, download_name='activity_log.csv')

@app.route('/download-log')
def download_log():
    if 'user_id' not in session:
        flash('You must log in first.', 'error')
        return redirect(url_for('login'))

    user_id = session['user_id']
    conn, cursor = get_db_connection()

    # ✅ Fetch user activity logs
    cursor.execute("SELECT timestamp, description FROM activity_logs WHERE user_id = %s", (user_id,))
    activities = cursor.fetchall()

    # ✅ Fetch risks
    cursor.execute("SELECT risk_type, severity, impact FROM risks WHERE user_id = %s", (user_id,))
    risks = cursor.fetchall()

    # ✅ Fetch threats
    cursor.execute("SELECT threat_name, threat_level, created_at FROM threats WHERE user_id = %s", (user_id,))
    threats = cursor.fetchall()

    # ✅ Fetch assets
    cursor.execute("SELECT asset_name, critical FROM assets WHERE user_id = %s", (user_id,))
    assets = cursor.fetchall()

    # ✅ Fetch mitigation planning
    cursor.execute("SELECT mitigation_name, status, last_updated FROM mitigation_plans WHERE user_id = %s", (user_id,))
    mitigations = cursor.fetchall()

    cursor.close()
    conn.close()

    # Convert all data to CSV format
    output = io.StringIO()
    csv_writer = csv.writer(output)

    # ✅ Write Headers
    csv_writer.writerow(["Timestamp", "Description"])
    for activity in activities:
        csv_writer.writerow([activity["timestamp"], activity["description"]])

    csv_writer.writerow([])  # Blank line for separation

    csv_writer.writerow(["Risk Type", "Severity", "Impact"])
    for risk in risks:
        csv_writer.writerow([risk["risk_type"], risk["severity"], risk["impact"]])

    csv_writer.writerow([])

    csv_writer.writerow(["Threat Name", "Threat Level", "Created At"])
    for threat in threats:
        csv_writer.writerow([threat["threat_name"], threat["threat_level"], threat["created_at"]])

    csv_writer.writerow([])

    csv_writer.writerow(["Asset Name", "Critical"])
    for asset in assets:
        csv_writer.writerow([asset["asset_name"], "Yes" if asset["critical"] else "No"])

    csv_writer.writerow([])

    csv_writer.writerow(["Mitigation Plan", "Status", "Last Updated"])
    for mitigation in mitigations:
        csv_writer.writerow([mitigation["mitigation_name"], mitigation["status"], mitigation["last_updated"]])

    output.seek(0)

    return send_file(
        io.BytesIO(output.getvalue().encode()),
        mimetype="text/csv",
        as_attachment=True,
        download_name="security_log.csv"
    )

    
@app.route('/signin', methods=['GET', 'POST'])
def signin():
    if request.method == 'POST':
        try:
            username = request.form['username']  # Capture the username
            fullname = request.form['fullname']
            email = request.form['email']
            password = request.form['password']
        except KeyError as e:
            flash(f'Missing form field: {e}', 'error')
            return redirect(url_for('signin'))

        # Hash the password before saving it
        hashed_password = generate_password_hash(password)

        conn, cursor = get_db_connection()
        
        # Check if the user already exists
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        existing_user = cursor.fetchone()
        
        if existing_user:
            flash('Email already registered. Please log in.', 'warning')
            return redirect(url_for('login'))
        
        # Insert new user into the database, with hashed password
        cursor.execute("INSERT INTO users (username, fullname, email, password_hash) VALUES (%s, %s, %s, %s)", 
                       (username, fullname, email, hashed_password))
        conn.commit()
        
        cursor.close()
        conn.close()
        
        flash('Account created successfully. Please log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('signin.html')

@app.route('/asset-management', methods=['GET', 'POST'])
def asset_management():
    conn, cursor = get_db_connection()
    
    if request.method == 'POST':
        try:
            asset_name = request.form['asset_name']
            asset_type = request.form['asset_type']
            description = request.form['description']
            criticality_level = request.form['criticality_level']
            owner = request.form['owner']

            # MODIFY THIS LINE (original was 'asset')
            cursor.execute(
                "INSERT INTO assets (asset_name, asset_type, description, criticality_level, owner) "  # Changed
                "VALUES (%s, %s, %s, %s, %s)",
                (asset_name, asset_type, description, criticality_level, owner)
            )
            conn.commit()
            flash('Asset saved successfully!', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'Error saving asset: {str(e)}', 'error')
    
    # MODIFY THIS LINE TOO (original was 'asset')
    cursor.execute("SELECT * FROM assets ORDER BY created_at DESC")  # Changed
    assets = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('AssetManagement.html', assets=assets)

@app.route('/mitigation-planning', methods=['GET', 'POST'])
def mitigation_planning():
    conn, cursor = get_db_connection()
    
    if request.method == 'POST':
        try:
            risk_id = request.form['risk_id']
            control_name = request.form['control_name']
            control_description = request.form['control_description']
            assigned_to = request.form['assigned_to']
            deadline = request.form['deadline']

            cursor.execute(
                "INSERT INTO mitigation_plans (risk_id, control_name, control_description, assigned_to, deadline) "
                "VALUES (%s, %s, %s, %s, %s)",
                (risk_id, control_name, control_description, assigned_to, deadline)
            )
            conn.commit()
            flash('Mitigation plan saved successfully!', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'Error saving mitigation plan: {str(e)}', 'error')
    
    # Fetch existing mitigation plans for display
    cursor.execute("SELECT * FROM mitigation_plans ORDER BY created_at DESC")
    mitigation_plans = cursor.fetchall()
    
    cursor.close()
    conn.close()

    return render_template('MitigationPlanning.html', mitigation_plans=mitigation_plans)

# Generate PDF
def generate_pdf(content, filename):
    try:
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer)
        text = p.beginText(40, 800)
        text.setFont("Helvetica", 12)
        for line in content.split("\n"):
            text.textLine(line)
        p.drawText(text)
        p.showPage()
        p.save()
        
        with open(filename, "wb") as f:
            f.write(buffer.getvalue())
        buffer.close()
    except Exception as e:
        logging.error(f"Error generating PDF: {e}")

# Generate DOCX
def generate_word(content, filename):
    try:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(filename)
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")

# Generate report file based on format
def generate_report_file(report_id, content, file_format):
    folder_path = "reports/"
    os.makedirs(folder_path, exist_ok=True)

    file_name = f"report_{report_id}.{file_format}"
    file_path = os.path.join(folder_path, file_name)

    if file_format == "pdf":
        pdfkit.from_string(content, file_path)
    elif file_format == "docx":
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_path)
    elif file_format == "xlsx":
        df = pd.DataFrame([line.split("|") for line in content.split("\n") if line])
        df.to_excel(file_path, index=False)

    return file_path


# View Reports Page
@app.route('/reports')
def reports():
    assets, risks, threats = fetch_data()
    conn, cursor = get_db_connection()
    try:
        cursor.execute("""
            SELECT r.id, r.title, r.created_at, r.format, r.status, r.file_path, u.username 
            FROM reports r
            JOIN users u ON r.user_id = u.id
            ORDER BY r.created_at DESC
        """)
        reports = cursor.fetchall()
    except Exception as e:
        flash(f'Error fetching reports: {str(e)}', 'error')
        reports = []
    finally:
        cursor.close()
        conn.close()
    return render_template('reports.html', assets=assets, risks=risks, threats=threats)

# Generate Report
@app.route('/generate_report', methods=['POST'])
def generate_report():
    assets, risks, threats = fetch_data()
    file_type = request.form.get('file_type')

    # Render the reports.html file
    html_content = render_template('reports.html', assets=assets, risks=risks, threats=threats)

    if file_type == "PDF":
        # Define PDF file path
        pdf_file_path = "static/report.pdf"

        # PDFKit Options
        options = {
            "enable-local-file-access": ""  # Fix ProtocolUnknownError
        }

        # Generate PDF
        pdfkit.from_string(html_content, pdf_file_path, configuration=PDFKIT_CONFIG, options=options)
        return send_file(pdf_file_path, as_attachment=True)

    elif file_type == "Word":
        doc_path = "static/report.docx"
        doc = Document()
        doc.add_heading("Cybersecurity Risk Assessment Report", level=1)

        doc.add_heading("Assets", level=2)
        for asset in assets:
            doc.add_paragraph(f"{asset[1]} ({asset[2]}) - {asset[3]}")

        doc.add_heading("Risks", level=2)
        for risk in risks:
            doc.add_paragraph(f"{risk[1]} ({risk[2]}) - {risk[3]}")

        doc.add_heading("Threats", level=2)
        for threat in threats:
            doc.add_paragraph(f"{threat[1]} ({threat[2]}) - {threat[3]}")

        doc.save(doc_path)
        return send_file(doc_path, as_attachment=True)

    return "Invalid format", 400



# Download Report
@app.route('/download-report/<int:report_id>')
def download_report(report_id):
    conn, cursor = get_db_connection()
    try:
        cursor.execute("SELECT file_path FROM reports WHERE id = %s", (report_id,))
        report = cursor.fetchone()

        if not report or not report["file_path"]:
            flash("Report not found", "error")
            return redirect(url_for("reports"))

        file_path = report["file_path"]

        # Check if the file exists before sending
        if not os.path.exists(file_path):
            flash("Report file is missing!", "error")
            return redirect(url_for("reports"))

        return send_file(file_path, as_attachment=True)
    except Exception as e:
        flash(f"Error downloading report: {str(e)}", "error")
        return redirect(url_for("reports"))
    finally:
        cursor.close()
        conn.close()





@app.route('/risk-assessment', methods=['GET', 'POST'])
def risk_assessment():
    if request.method == 'POST':
        conn, cursor = get_db_connection()
        try:
            # Existing assessment save
            asset = request.form['asset']
            threat = request.form['threat']
            likelihood = request.form['likelihood']
            impact = request.form['impact']

            # Save assessment
            cursor.execute(
                "INSERT INTO risk_assessments (asset, threat, likelihood, impact) "
                "VALUES (%s, %s, %s, %s) RETURNING id",
                (asset, threat, likelihood, impact)
            )
            assessment_id = cursor.fetchone()[0]

            # Calculate risk type
            risk_type = calculate_risk_type(likelihood, impact)

            # Create corresponding risk entry
            cursor.execute(
                "INSERT INTO risks (assessment_id, affected_asset, threat_name, risk_type) "
                "VALUES (%s, %s, %s, %s)",
                (assessment_id, asset, threat, risk_type)
            )

            conn.commit()
            flash('Risk assessment saved and monitoring entry created!', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'Error: {str(e)}', 'error')
        finally:
            cursor.close()
            conn.close()
        return redirect(url_for('risk_assessment'))

    # GET request handling
    conn, cursor = get_db_connection()
    cursor.execute("SELECT * FROM risk_assessments ORDER BY created_at DESC")
    assessments = cursor.fetchall()
    cursor.close()
    conn.close()
    
    return render_template('RiskAssessment.html', assessments=assessments)

def calculate_risk_type(likelihood, impact):
    likelihood_weight = {'Low': 1, 'Medium': 2, 'High': 3}
    impact_weight = {'Low': 1, 'Medium': 2, 'High': 3}
    score = likelihood_weight[likelihood] * impact_weight[impact]
    
    if score >= 6: return 'High Risk'
    elif score >= 3: return 'Medium Risk'
    return 'Low Risk'
@app.route('/risk-monitoring')
def risk_monitoring():
    conn, cursor = get_db_connection()
    
    # Get all risks including assessments
    cursor.execute("""
        SELECT r.id, r.file_id, r.affected_asset, r.threat_name,
               COALESCE(to_char(r.detected_time, 'HH12:MI AM'), 'Assessment') as time,
               r.file_size, r.risk_type, ra.likelihood, ra.impact
        FROM risks r
        LEFT JOIN risk_assessments ra ON r.assessment_id = ra.id
        ORDER BY COALESCE(r.detected_time, ra.created_at) DESC
    """)
    risks = cursor.fetchall()
    
    # Statistics
    cursor.execute("""
        SELECT 
            COUNT(*) FILTER (WHERE risk_type = 'High Risk') as high_risk,
            COUNT(*) FILTER (WHERE risk_type = 'Medium Risk') as medium_risk,
            COUNT(*) FILTER (WHERE risk_type = 'Low Risk') as low_risk
        FROM risks
    """)
    risk_stats = cursor.fetchone()
    
    cursor.close()
    conn.close()

    return render_template('RiskMonitoring.html', risks=risks, risk_stats=risk_stats)

@app.route('/settings')
def settings():
    return render_template('Settings.html')

@app.route('/threat-management', methods=['GET', 'POST'])
def threat_management():
    conn, cursor = get_db_connection()
    
    # Handle form submission
    if request.method == 'POST':
        threat_name = request.form.get('threat-name')
        threat_source = request.form.get('threat-source')
        description = request.form.get('threat-description')
        severity = request.form.get('threat-severity')
        
        try:
            cursor.execute(
                "INSERT INTO threats (threat_name, threat_source, description, severity) "
                "VALUES (%s, %s, %s, %s)",
                (threat_name, threat_source, description, severity)
            )
            conn.commit()
            flash('Threat added successfully!', 'success')
        except Exception as e:
            conn.rollback()
            flash(f'Error saving threat: {str(e)}', 'error')
    
    # Get threats for real-time feed
    cursor.execute("SELECT * FROM threats ORDER BY created_at DESC")
    threats = cursor.fetchall()
    cursor.close()
    conn.close()
    
    print("Fetched threats:", threats)  # Debugging line

    return render_template('ThreatManagement.html', threats=threats)

@app.route('/threat-details/<int:threat_id>')
def threat_details(threat_id):
    conn, cursor = get_db_connection()
    cursor.execute("SELECT * FROM threats WHERE id = %s", (threat_id,))
    threat = cursor.fetchone()
    cursor.close()
    conn.close()

    print("Fetched threat details for ID", threat_id, ":", threat)  # Debugging line

    if threat:
        return jsonify({
            'threat_name': threat['threat_name'],
            'threat_source': threat['threat_source'],
            'description': threat['description'],
            'severity': threat['severity'],
            'created_at': threat['created_at'].strftime('%Y-%m-%d %H:%M:%S UTC')
        })
    else:
        return jsonify({'error': 'Threat not found'}), 404
if __name__ == '__main__':
    app.run(debug=True)